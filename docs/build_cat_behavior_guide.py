from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "cat-behavior-translation-guide-zh.docx"

# compact_reference_guide token map
FONT = "Hiragino Sans GB"
CJK_FONT = "Hiragino Sans GB"
INK = "211D1A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
ORANGE = "D94F2B"  # named cover/caution override, reused consistently
GOLD = "A16C00"    # named evidence-label override
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 100
CELL_SIDE = 120

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)


def set_run_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_style_font(style, size, bold=False, color=INK):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


styles = doc.styles
normal = styles["Normal"]
set_style_font(normal, 11, False, INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Title", 30, INK, 0, 8),
    ("Subtitle", 14, DARK_BLUE, 0, 16),
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[name]
    set_style_font(style, size, name != "Subtitle", color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number"]:
    style = styles[name]
    set_style_font(style, 11, False, INK)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

if "Small Meta" not in styles:
    small_meta = styles.add_style("Small Meta", WD_STYLE_TYPE.PARAGRAPH)
else:
    small_meta = styles["Small Meta"]
set_style_font(small_meta, 8.5, False, MUTED)
small_meta.paragraph_format.space_after = Pt(4)
small_meta.paragraph_format.line_spacing = 1.1

if "Table Citation" not in styles:
    citation_style = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
else:
    citation_style = styles["Table Citation"]
set_style_font(citation_style, 8.5, False, MUTED)
citation_style.paragraph_format.space_before = Pt(4)
citation_style.paragraph_format.space_after = Pt(4)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_SIDE, bottom=CELL_TOP_BOTTOM, end=CELL_SIDE):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, size=9.4, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), size=size, bold=bold, color=color)


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[idx], header, size=9, bold=True, color=DARK_BLUE,
                      align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for row_data in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value, size=9.1,
                          align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    doc.add_paragraph("", style="Table Citation")
    return table


def add_para(text="", style=None, size=None, bold=None, color=INK, italic=None,
             align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size or 11, bold=bold, color=color, italic=italic)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_run_font(p.add_run(text), 11)
    return p


def add_number(text):
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run(text), 11)
    return p


def add_callout(label, text, fill=LIGHT_GRAY, accent=BLUE):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "AFAFAF")
        borders.append(border)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(label + "  "), 10, bold=True, color=accent)
    set_run_font(p.add_run(text), 10.5, color=INK)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run("貓語行為觀察室  |  研究與照護指南"), 8.5, bold=True, color=MUTED)
    p.paragraph_format.space_after = Pt(0)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("教育用途 · 不取代獸醫診斷"), 8, color=MUTED)


set_header_footer(section)

# Editorial cover using the compact reference guide with named orange/gold cover overrides.
add_para("EVIDENCE-INFORMED FIELD GUIDE", size=9, bold=True, color=ORANGE,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=18)
add_para("貓咪行為翻譯指南", size=30, bold=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para("從姿勢、耳朵、尾巴、眼睛、聲音與情境，判讀較可能的情緒狀態",
         style="Subtitle", align=WD_ALIGN_PARAGRAPH.CENTER, after=32)
add_para("不是貓語字典，而是一套以研究與臨床指南為基礎的觀察方法",
         size=11, color=DARK_BLUE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=104)
add_para("研究整理日期", size=9, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add_para("2026 年 7 月 20 日", size=12, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para("適用對象：一般貓咪照護者與內容設計參考", size=9.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

doc.add_page_break()

doc.add_heading("先說結論：這不是逐字翻譯", level=1)
add_callout(
    "核心原則",
    "貓咪的單一動作通常沒有唯一意思。可靠判讀要同時看「訊號組合、發生情境、個體平常基準、行為是否突然改變」。",
    fill="FFF4E8", accent=ORANGE,
)
add_para(
    "現有研究能協助我們辨識正向社交、害怕／焦慮、高度喚起、疼痛風險與資源衝突等狀態，"
    "但不能把一聲喵或一個尾巴動作固定翻成一句人話。這份指南刻意使用「較可能」「需要搭配」與「需排除」等表述，"
    "避免把機率性線索說成確定答案。[1–6]"
)

doc.add_heading("四步判讀法", level=2)
for item in [
    "看整體：先看身體是柔軟還是緊繃，再看耳、尾、眼、鬍鬚與聲音。",
    "看組合：至少兩到三個方向一致的訊號，才提高判讀信心。",
    "看情境：光線會影響瞳孔；觸摸、陌生刺激、食物期待與疼痛會改變同一動作的含義。",
    "看變化：和平常基準相比，突然躲藏、不吃、亂尿、不願跳或攻擊，應先考慮健康與壓力。",
]:
    add_number(item)

doc.add_heading("證據等級", level=2)
add_table(
    ["等級", "使用方式", "代表資料"],
    [
        ("A", "可作為較強提示，但仍不能單獨診斷。", "經驗證量表、控制實驗或明確行為研究，例如 Feline Grimace Scale、慢眨眼實驗。"),
        ("B", "適合形成照護建議與情境判讀。", "同儕審查觀察研究、專業臨床指南與獸醫機構資料。"),
        ("C", "只作為探索提示，需特別依賴個體基準。", "常見照護解讀、研究仍有限或高度依情境而變的動作。"),
    ],
    [900, 2900, 5560],
)

doc.add_page_break()
doc.add_heading("快速辨識：放鬆與正向社交", level=1)
add_para("以下訊號在彼此一致、身體柔軟且貓可自由離開時，較支持安全與正向互動。[2–4, 7]")
add_table(
    ["訊號", "較可能代表", "照護者怎麼做", "限制／反例"],
    [
        ("尾巴直立並靠近", "社交接近、願意互動（B）", "停下讓牠先嗅聞；讓牠決定距離。", "不是固定的「開心」標誌；社會關係與情境會影響使用。"),
        ("慢慢眨眼", "正向情緒溝通、較願意接近（A）", "可慢眨回應並稍微移開凝視。", "單側瞇眼、分泌物或持續閉眼需排除眼部不適。"),
        ("耳朵自然向前", "放鬆或對環境有興趣（B）", "保持緩慢、可預測的互動。", "若身體緊、瞳孔大，可能是高度注意而非放鬆。"),
        ("身體柔軟、伸展", "較安全、休息或舒適（B）", "安靜陪伴；不要因露肚就直接摸腹部。", "腹部暴露若伴隨緊繃、四肢準備踢，可能是防禦。"),
        ("臉頰／身體磨蹭", "氣味標記，也常見於熟悉的社交互動（B）", "讓牠主動磨蹭；可輕撫頭頸。", "反覆、急促磨蹭並伴隨踱步或大叫，可能是挫折。"),
        ("咕嚕＋全身放鬆", "常見於舒適、接觸維持或要求互動（B）", "以整體姿勢與情境確認。", "疼痛、緊張與臨終狀態也可能咕嚕；不能單獨視為開心。"),
    ],
    [1250, 2200, 2600, 3310],
)

doc.add_page_break()

doc.add_heading("快速辨識：好奇、狩獵與玩耍", level=1)
add_table(
    ["訊號", "較可能代表", "照護者怎麼做", "限制／反例"],
    [
        ("耳朵向前、鬍鬚前伸", "注意力朝向刺激、探索或狩獵準備（B）", "讓牠觀察；提供距離與退路。", "若身體壓低且僵硬，可能是恐懼或衝突。"),
        ("瞳孔放大", "喚起程度提高；可能是玩耍、驚訝、恐懼或疼痛（B）", "先確認光線，再看全身與觸發事件。", "昏暗環境自然會放大瞳孔，不能單獨判讀情緒。"),
        ("盯住、屁股扭動、撲擊玩具", "狩獵序列與遊戲（B）", "用逗貓棒模擬獵物，讓牠成功捕捉後結束。", "不要用手腳當玩具；對另一隻貓的追逐需確認是否互惠。"),
        ("對窗外快速顫顎／喀喀聲", "常見於看到獵物時的高度注意（C）", "提供安全觀察位置與室內遊戲。", "功能尚未有單一確定結論，不宜翻成固定語句。"),
        ("互惠摔抱、少叫聲", "較接近社交玩耍（A/B）", "確認角色輪替、身體鬆、結束後可恢復平靜。", "若只有一方追、堵路或持續叫喊，可能已變成衝突。"),
    ],
    [1250, 2200, 2600, 3310],
)

doc.add_heading("玩耍還是打架？", level=2)
add_callout(
    "研究提示",
    "一項分析 105 組雙貓互動的研究發現，互惠摔抱最接近玩耍；叫聲與追逐較接近對抗。研究也辨識出「中間型」互動，說明並非所有片段都能二分為玩或打。[6]",
)
for item in [
    "較像玩耍：角色交換、動作有停頓、身體相對放鬆、咬合受抑制、叫聲少。",
    "較像衝突：非互惠追逐、堵住資源或通道、凝視、低吼／嘶叫、尾巴膨起、互動後仍躲避。",
    "安全做法：不要徒手拉開；以屏障或遠距離方式中斷，讓雙方分開冷靜並檢查受傷。",
]:
    add_bullet(item)

doc.add_heading("快速辨識：擔心、害怕與防禦", level=1)
add_table(
    ["訊號", "較可能代表", "照護者怎麼做", "避免做什麼"],
    [
        ("身體壓低蜷縮", "降低能見度、準備逃避，常見於擔心或害怕（B）", "停止靠近，提供可躲藏的箱子、高處或外出籠。", "不要把牠從躲藏處拖出來。"),
        ("尾巴貼身或捲住身體", "緊張、試圖保護身體（B）", "降低聲音、光線與人流。", "不要強抱或逼牠面對刺激。"),
        ("耳朵轉向側後／壓平", "警戒、害怕或防禦升級（B）", "拉開距離，等待耳位與肌肉放鬆。", "不要凝視、伸手靠近臉部。"),
        ("拱背、側身、毛豎起", "使體型看起來更大，常見於強烈恐懼或防禦（B）", "給出清楚退路並遠離。", "不要追逐、處罰或大聲喝斥。"),
        ("嘶聲、低吼、抬起前掌", "明確的距離增加訊號，可能即將揮擊（B）", "立刻停止互動，讓牠退開。", "不要把嘶聲解讀成壞脾氣。"),
        ("躲藏", "壓力下的正常因應方式，也可能是疼痛／疾病線索（B）", "保留安全處並觀察食慾、排泄與日常活動。", "若突然增加或伴隨生理變化，不要只當個性問題。"),
    ],
    [1250, 2320, 2820, 2970],
)

doc.add_page_break()

doc.add_heading("快速辨識：過度刺激、挫折與衝突", level=1)
add_table(
    ["訊號", "較可能代表", "建議回應", "注意事項"],
    [
        ("尾巴快速、大幅拍打", "喚起升高、挫折或觸摸耐受度降低（B）", "立即暫停撫摸，讓牠自行離開。", "細小尾尖移動可能只是專注；速度和全身張力很重要。"),
        ("皮膚抽動、耳朵側轉", "接觸刺激累積、注意力轉向（B）", "把手移開，不要測試「還能摸多久」。", "若頻繁發生或伴隨疼痛，請諮詢獸醫。"),
        ("突然回頭咬／抓", "可能是過度刺激、疼痛、恐懼或低接觸耐受（B）", "記錄觸摸部位與前兆；改為短暫、由貓主動的互動。", "不是毫無預警；常有細微前兆。"),
        ("踱步、急促磨蹭、大叫", "挫折或高度期待，也可能是發情、甲狀腺或認知問題（B）", "檢查需求與環境；突然出現或持續時就醫。", "不要只用食物或注意力壓下行為而忽略原因。"),
        ("噴尿：站立、尾巴顫動、對垂直面少量排尿", "氣味標記；可能受領域威脅、壓力或性荷爾蒙影響（B）", "先做醫療檢查，再改善資源分布與外來貓刺激。", "不要處罰；處罰可能增加焦慮。"),
    ],
    [1400, 2400, 2800, 2760],
)

doc.add_page_break()
doc.add_heading("正常行為，不等於不需要理解", level=1)
add_table(
    ["行為", "功能／較可能含義", "何時值得關注"],
    [
        ("抓抓", "天生的伸展、磨除爪鞘、視覺與氣味標記行為。[8]", "抓點突然改變、減少或只抓低處，可能與活動力或疼痛相關。"),
        ("磨蹭", "在物體或熟悉對象留下臉部／身體氣味，也可出現在正向社交。[1, 8]", "若伴隨急躁、踱步、叫聲或其他異常，需看整體情境。"),
        ("踩奶", "常見於舒適、熟悉接觸或氣味標記情境；個體差異大（C）。", "若伴隨啃食布料、皮膚受傷或難以中斷，請諮詢專業人士。"),
        ("舔毛", "日常整理與氣味維持；也可能是轉移／自我安撫行為。[9]", "脫毛、紅腫、破皮、局部集中舔或突然大增，需要先排除搔癢與疼痛。"),
        ("躲藏", "面對威脅時的因應策略；安全躲藏處可提升控制感。[1]", "突然長時間躲藏，尤其伴隨不吃、活動下降或排泄改變。"),
        ("亂尿／排便", "可能是砂盆偏好、資源衝突、壓力或醫療問題；不是報復。[10]", "任何突然改變都應先做獸醫檢查，尤其頻繁進出砂盆卻排不出尿。"),
    ],
    [1250, 4550, 3560],
)

doc.add_page_break()

doc.add_heading("疼痛與健康警訊", level=1)
add_callout(
    "重要",
    "貓常會隱藏虛弱與疼痛。行為改變可能比明顯叫痛更早出現；網站或本指南不能取代臨床檢查。[1, 5]",
    fill="FDEDE8", accent=ORANGE,
)

doc.add_heading("Feline Grimace Scale 的五個臉部行動單元", level=2)
add_para(
    "Feline Grimace Scale（FGS）在自然發生急性疼痛的貓中完成建立與驗證。五項線索各評 0–2 分："
    "耳位、眼眶收緊、口鼻張力、鬍鬚位置與頭部位置。原始研究中，FGS 與另一個急性疼痛量表高度相關，"
    "也能反映止痛後的變化。[5]"
)
add_table(
    ["臉部線索", "疼痛相關變化", "觀察提醒"],
    [
        ("耳朵", "耳尖分開、向外旋轉或壓低", "需要和眼、口鼻、鬍鬚及頭位一起評估。"),
        ("眼睛", "眼眶收緊、瞇眼或緊閉", "單側瞇眼也可能是眼部疾病；需就醫。"),
        ("口鼻", "從圓潤變得緊繃、扁平或橢圓", "品種頭型可能影響外觀。"),
        ("鬍鬚", "變直、前伸、呈刺狀", "興奮／狩獵時也可能前伸，需搭配其他單元。"),
        ("頭位", "頭低於肩線或下垂", "疲倦、鎮靜或環境也可能影響姿勢。"),
    ],
    [1400, 4300, 3660],
)
add_para(
    "FGS 是臨床疼痛評估工具，不應由網站把幾個勾選結果直接轉成診斷。原始研究的介入門檻與計分方式應由獸醫或經訓練人員使用。",
    style="Table Citation",
)

doc.add_heading("何時應諮詢獸醫", level=2)
for item in [
    "行為與平常不同且持續：躲藏、攻擊、活動或互動大幅下降。",
    "食慾下降、體重改變、嘔吐／腹瀉、飲水或排尿量改變。",
    "不願跳躍、跛行、碰觸特定位置會躲或攻擊。",
    "過度舔毛、脫毛、皮膚紅腫或傷口。",
    "在砂盆外排泄、頻繁進出砂盆、排尿用力或叫聲。",
    "臉部疼痛表情持續，或與呼吸、活動與食慾異常同時出現。",
]:
    add_bullet(item)

add_callout(
    "緊急警訊",
    "呼吸費力或張口呼吸、無法排尿、昏厥、癲癇、嚴重外傷、持續嘔吐、明顯虛弱或意識改變，應立即尋求急診獸醫協助。",
    fill="FDEDE8", accent=ORANGE,
)

doc.add_heading("把研究轉成網站判讀規則", level=1)
add_para(
    "網站採用加權線索群組，而不是「一個按鈕＝一句翻譯」。使用者各選一項身體姿勢、耳位、尾巴、臉部與情境；"
    "系統把線索彙整成五種可能狀態，並顯示判讀信心。"
)
add_table(
    ["可能狀態", "主要線索群", "網站輸出原則"],
    [
        ("放鬆／正向社交", "身體柔軟、耳朵自然、慢眨眼、直尾接近", "鼓勵由貓主動、可隨時結束的互動。"),
        ("好奇／注意提高", "耳與鬍鬚向前、瞳孔變化、探索姿勢", "先排除光線；維持距離與退路。"),
        ("害怕／焦慮", "壓低蜷縮、尾巴貼身、耳朵後轉／壓平", "停止逼近，提供躲藏與可預測環境。"),
        ("高度喚起／過度刺激", "快速甩尾、肌肉緊、壓耳、叫聲或非互惠追逐", "停止觸摸或衝突，避免徒手介入。"),
        ("疼痛風險", "FGS 臉部線索、突然行為改變、活動／食慾／排泄異常", "明確標示不能診斷，建議記錄並諮詢獸醫。"),
    ],
    [1850, 3910, 3600],
)

doc.add_heading("判讀限制", level=2)
for item in [
    "不同個體、品種、年齡、早期社會化與生活經驗會改變訊號表現。",
    "照片會失去時間序列與情境；短片通常比單張照片更有價值。",
    "瞳孔受光線影響；耳、鬍鬚與臉型也可能受品種外觀影響。",
    "同一行為可能同時服務多個功能，例如咕嚕、磨蹭、舔毛與躲藏。",
    "線上工具最多提供教育性提示，不能排除疼痛、神經、皮膚、內分泌或泌尿疾病。",
]:
    add_bullet(item)

doc.add_heading("參考資料", level=1)
references = [
    ("[1]", "Ellis SLH, Rodan I, Carney HC, et al. AAFP and ISFM Feline Environmental Needs Guidelines. Journal of Feline Medicine and Surgery. 2013;15(3):219–230.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC11383066/"),
    ("[2]", "Humphrey T, Proops L, Forman J, Spooner R, McComb K. The role of cat eye narrowing movements in cat–human communication. Scientific Reports. 2020;10:16503.",
     "https://pubmed.ncbi.nlm.nih.gov/33020542/"),
    ("[3]", "Deputte BL, Doll A. Heads and Tails: An Analysis of Visual Signals in Cats, Felis catus. Animals. 2021;11(9):2752.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC8469685/"),
    ("[4]", "Cats Protection. Cat Body Language. 專業照護資訊頁面。",
     "https://www.cats.org.uk/help-and-advice/cat-behaviour/cat-body-language"),
    ("[5]", "Evangelista MC, Watanabe R, Leung VSY, et al. Facial expressions of pain in cats: the development and validation of a Feline Grimace Scale. Scientific Reports. 2019;9:19128.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC6911058/"),
    ("[6]", "Gajdoš-Kmecová N, Peták I, Kottferová J, et al. An ethological analysis of close-contact inter-cat interactions determining if cats are playing, fighting, or something in between. Scientific Reports. 2023;13:92.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC9879969/"),
    ("[7]", "RSPCA. Understanding Your Cat’s Behaviour／Cat Body Language. 動物福利照護指南。",
     "https://www.rspca.org.uk/adviceandwelfare/pets/cats/behaviour"),
    ("[8]", "Cornell Feline Health Center. Feline Behavior Problems: Destructive Behavior／House Soiling.",
     "https://www.vet.cornell.edu/departments-centers-and-institutes/cornell-feline-health-center/health-information/feline-health-topics/feline-behavior-problems-destructive-behavior"),
    ("[9]", "Cornell Feline Health Center. Cats that Lick Too Much.",
     "https://www.vet.cornell.edu/departments-centers-and-institutes/cornell-feline-health-center/health-information/feline-health-topics/cats-lick-too-much"),
    ("[10]", "Carney HC, Sadek TP, Curtis TM, et al. AAFP and ISFM Guidelines for Diagnosing and Solving House-Soiling Behavior in Cats. Journal of Feline Medicine and Surgery.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC11104062/"),
    ("[11]", "Tavernier C, Ahmed S, Houpt KA, Yeon SC. Feline vocal communication. Journal of Veterinary Science. 2020;21(1):e18.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC7000907/"),
    ("[12]", "2024 AAFP Intercat Tension Guidelines: Recognition, Prevention and Management. Journal of Feline Medicine and Surgery.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC11292941/"),
]
for label, citation, url in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(label + " "), 9.5, bold=True, color=BLUE)
    set_run_font(p.add_run(citation + " "), 9.5, color=INK)
    set_run_font(p.add_run(url), 8.5, color=DARK_BLUE)

doc.add_heading("使用聲明", level=2)
add_para(
    "本指南為教育與產品設計用途，整理公開研究與專業機構資料；不是診斷、治療或個別行為諮詢。"
    "若貓咪出現突然或持續的行為改變、疼痛線索或緊急症狀，請聯絡合格獸醫師。資料會隨研究進展更新。"
)

# Prevent orphan headings and normalize all table cells one last time.
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
